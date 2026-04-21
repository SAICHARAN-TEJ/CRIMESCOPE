"""
CrimeScope — Document Agent (Antigravity-Hardened + Guardian Pattern).

Parses PDF and DOCX files into structured text chunks.

v4.2 Hardening:
  - Guardian Input: Encrypted PDF detection, corrupted DOCX handling
  - Guardian Output: Ensures text extraction is non-empty
  - Chaos injection: Random failures under chaos mode
  - Mixed content handling: Images in PDFs logged, OCR skipped gracefully
  - Null-byte / control char stripping in extracted text
"""

from __future__ import annotations

import io
import re
from typing import Any

from app.core.config import get_settings
from app.core.logger import get_logger
from app.engine.agents.base import BaseAgent, DataIntegrityError, chaos_injector
from app.schemas.events import AgentResult, AgentType
from app.storage.minio_client import get_minio

logger = get_logger("crimescope.agent.document")

# ── Safety limits ─────────────────────────────────────────────────────────
MAX_DOC_SIZE_BYTES = 500 * 1024 * 1024  # 500MB hard cap
MAX_PAGES = 2000                         # Skip docs with absurd page counts
MIN_EXTRACTED_CHARS = 10                 # Below this, extraction is considered failed
ALLOWED_DOC_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".rtf", ".md", ".csv"}

# Control chars to strip from extracted text (keep newlines, tabs)
_CONTROL_CHAR_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")


def _clean_text(text: str) -> str:
    """Strip null bytes, control characters, and excessive whitespace."""
    text = _CONTROL_CHAR_RE.sub("", text)
    # Collapse runs of 3+ newlines to 2
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _detect_encrypted_pdf(data: bytes) -> bool:
    """Quick check for encrypted PDFs by scanning header bytes."""
    try:
        header = data[:2048].decode("latin-1", errors="ignore")
        return "/Encrypt" in header
    except Exception:
        return False


def _parse_pdf(data: bytes) -> tuple[str, list[str]]:
    """
    Extract text from PDF bytes.
    Returns (text, warnings).
    """
    warnings: list[str] = []

    # ── Guard: Encrypted PDF ──────────────────────────────────────
    if _detect_encrypted_pdf(data):
        warnings.append("⚠ PDF is encrypted — text extraction may be incomplete")

    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data))

        # ── Guard: Excessive pages ────────────────────────────────
        if len(reader.pages) > MAX_PAGES:
            warnings.append(f"⚠ PDF has {len(reader.pages)} pages — truncating to {MAX_PAGES}")

        pages = []
        image_page_count = 0
        for i, page in enumerate(reader.pages[:MAX_PAGES]):
            try:
                text = page.extract_text() or ""
                pages.append(text)

                # Detect image-heavy pages (likely scanned docs)
                if len(text.strip()) < 20 and hasattr(page, "images"):
                    try:
                        if len(page.images) > 0:
                            image_page_count += 1
                    except Exception:
                        pass
            except Exception as e:
                warnings.append(f"⚠ Page {i + 1} extraction failed: {type(e).__name__}")
                pages.append("")

        if image_page_count > 0:
            warnings.append(
                f"⚠ {image_page_count} pages appear to be scanned images — "
                f"OCR not available, text may be incomplete"
            )

        return _clean_text("\n\n".join(pages)), warnings

    except Exception as e:
        warnings.append(f"PDF parse failed: {e}")
        return "", warnings


def _parse_docx(data: bytes) -> tuple[str, list[str]]:
    """
    Extract text from DOCX bytes.
    Returns (text, warnings).
    """
    warnings: list[str] = []
    try:
        from docx import Document
        doc = Document(io.BytesIO(data))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Also extract table content
        table_text = []
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    table_text.append(" | ".join(cells))

        if table_text:
            warnings.append(f"Extracted {len(table_text)} table rows")

        full_text = "\n".join(paragraphs)
        if table_text:
            full_text += "\n\n--- Tables ---\n" + "\n".join(table_text)

        return _clean_text(full_text), warnings
    except Exception as e:
        warnings.append(f"DOCX parse failed: {type(e).__name__}: {e}")
        # Try raw text fallback for corrupted DOCX
        try:
            raw = data.decode("utf-8", errors="replace")
            # Extract anything that looks like text from the binary
            text_parts = re.findall(r"[A-Za-z][A-Za-z\s.,;:!?'-]{20,}", raw)
            if text_parts:
                warnings.append("⚠ DOCX corrupted — extracted partial text via fallback")
                return _clean_text("\n".join(text_parts[:100])), warnings
        except Exception:
            pass
        return "", warnings


def _chunk_text(text: str, max_tokens: int = 1500, overlap: int = 200) -> list[str]:
    """Split text into overlapping chunks for LLM context windows."""
    words = text.split()
    if len(words) <= max_tokens:
        return [text] if text.strip() else []
    chunks = []
    i = 0
    while i < len(words):
        chunk = " ".join(words[i : i + max_tokens])
        chunks.append(chunk)
        i += max_tokens - overlap
    return chunks


class DocumentAgent(BaseAgent):
    """
    Parses uploaded documents (PDF, DOCX, TXT) into text chunks.
    Downloads from MinIO, extracts content, and returns structured facts.

    Guardian Pattern:
      - Input: Validates file list, rejects unsupported formats
      - Output: Ensures at least one chunk is extracted or a warning is produced
    """

    agent_type = AgentType.DOCUMENT
    agent_name = "document_agent"

    # ── Guardian: Input Validation ────────────────────────────────────

    def validate_input(self, job_id: str, payload: dict[str, Any]) -> None:
        """Validate document processing inputs."""
        super().validate_input(job_id, payload)
        files = payload.get("files")
        if files is not None and not isinstance(files, list):
            raise DataIntegrityError(self.agent_name, "files must be a list")

    # ── Guardian: Output Validation ───────────────────────────────────

    def validate_output(self, result: AgentResult) -> None:
        """Validate document processing outputs."""
        super().validate_output(result)
        # Must have at least one fact reporting results
        if result.success and not result.facts:
            raise DataIntegrityError(
                self.agent_name,
                "Document agent returned success but produced no facts",
                recoverable=False,
            )

    # ── Core Execution (with chaos injection) ─────────────────────────

    @chaos_injector
    async def _execute(self, job_id: str, payload: dict[str, Any]) -> AgentResult:
        files = payload.get("files", [])
        doc_files = [
            f for f in files
            if not f.get("content_type", "").startswith("video/")
        ]

        if not doc_files:
            return AgentResult(
                agent=self.agent_type,
                success=True,
                facts=["No document files to process"],
            )

        minio = get_minio()
        all_chunks: list[str] = []
        all_facts: list[str] = []

        for df in doc_files:
            object_key = df.get("object_key", "")
            filename = df.get("filename", "unknown")
            content_type = df.get("content_type", "")

            # ── Input Guard: Validate object_key exists ───────────────
            if not object_key:
                all_facts.append(f"⚠ Skipped {filename}: missing object_key")
                continue

            # ── Download file ─────────────────────────────────────────
            try:
                data = minio.get_object_bytes(object_key)
            except Exception as e:
                all_facts.append(f"⚠ Download failed for {filename}: {type(e).__name__}")
                continue

            if not data:
                all_facts.append(f"⚠ Could not download {filename}")
                continue

            # ── Input Guard: Size validation ──────────────────────────
            if len(data) > MAX_DOC_SIZE_BYTES:
                all_facts.append(
                    f"⚠ Skipped {filename}: {len(data) / 1e6:.1f}MB exceeds "
                    f"{MAX_DOC_SIZE_BYTES / 1e6:.0f}MB limit"
                )
                continue

            if len(data) == 0:
                all_facts.append(f"⚠ Skipped {filename}: 0-byte file")
                continue

            # ── Parse based on content type ───────────────────────────
            text = ""
            warnings: list[str] = []

            try:
                if "pdf" in content_type or filename.lower().endswith(".pdf"):
                    text, warnings = _parse_pdf(data)
                elif "docx" in content_type or filename.lower().endswith(".docx"):
                    text, warnings = _parse_docx(data)
                else:
                    text = _clean_text(data.decode("utf-8", errors="replace"))
            except Exception as e:
                all_facts.append(f"⚠ Parse error for {filename}: {type(e).__name__}: {e}")
                continue

            # Report any parser warnings
            all_facts.extend(warnings)

            # ── Output Guard: Check extraction result ─────────────────
            if not text.strip() or len(text.strip()) < MIN_EXTRACTED_CHARS:
                all_facts.append(f"⚠ No meaningful text extracted from {filename} ({len(text)} chars)")
                continue

            chunks = _chunk_text(text)
            all_chunks.extend(chunks)
            all_facts.append(
                f"Parsed {filename}: {len(text)} chars → {len(chunks)} chunks"
            )

        # Store chunks in payload for downstream agents
        payload["text_chunks"] = all_chunks

        return AgentResult(
            agent=self.agent_type,
            success=True,
            facts=all_facts if all_facts else ["No documents processed"],
        )
